import xml.etree.ElementTree as ET
import re
import os
import glob
from docx import Document
from docx.shared import Inches
import base64
import tempfile
from html.parser import HTMLParser
import urllib.parse

def extract_questions_and_answers(xml_string):
    root = ET.fromstring(xml_string)
    output = []
    
    for question in root.findall("question"):
        # Убираем HTML из текста вопроса
        question_text_elem = question.find("questiontext/text")
        question_text_raw = question_text_elem.text if question_text_elem is not None and question_text_elem.text is not None else ""
        question_text = clean_html(question_text_raw).strip()
        
        answers = question.findall("answer")
        all_answers = []
        for answer in answers:
            # Убираем HTML из текста ответа
            answer_text_elem = answer.find("text")
            answer_text_raw = answer_text_elem.text if answer_text_elem is not None and answer_text_elem.text is not None else ""
            answer_text = clean_html(answer_text_raw).strip()
            fraction = float(answer.attrib.get("fraction", "0"))
            if fraction > 0:
                all_answers.append(("+", answer_text))
            else:
                all_answers.append(("", answer_text))  # Неправильный ответ без подписи
        
        if all_answers:
            output.append({
                "question": question_text,  # очищенный текст для txt
                "answers": all_answers,
                "question_elem": question  # для доступа к <file>
            })
    
    return output

def clean_html(raw_html):
    """Функция для удаления HTML-тегов и преобразования формул."""
    # Заменяем <br> на перевод строки
    clean_text = re.sub(r'<br\s*/?>', '\n', raw_html, flags=re.IGNORECASE)
    # Удаляем все HTML-теги
    clean_text = re.sub(r'<[^>]+>', '', clean_text)
    # Заменяем &nbsp; на пробел
    clean_text = re.sub(r'&nbsp;', ' ', clean_text)
    # Заменяем &mdash; на -
    clean_text = re.sub(r'&mdash;', '-', clean_text)
    # Заменяем &lt; на <
    clean_text = re.sub(r'&lt;', '<', clean_text)
    # Заменяем &gt; на >
    clean_text = re.sub(r'&gt;', '>', clean_text)
    # Преобразуем формулы LaTeX вида \( ... \) в обычный текст без спецсимволов
    def latex_to_text(match):
        formula = match.group(1)
        # Заменяем основные LaTeX-команды на символы
        latex_replacements = {
            r'\\leq': '≤',
            r'\\geq': '≥',
            r'\\cdot': '*',
            r'\\times': '×',
            r'\\pm': '±',
            r'\\div': '÷',
            r'\\neq': '≠',
            r'\\approx': '≈',
            r'\\sim': '∼',
            r'\\infty': '∞',
            r'\\rightarrow': '→',
            r'\\leftarrow': '←',
            r'\\Rightarrow': '⇒',
            r'\\Leftarrow': '⇐',
            r'\\to': '→',
            r'\\sqrt': '√',
        }
        for latex, symbol in latex_replacements.items():
            formula = re.sub(latex, symbol, formula)
        # Удаляем фигурные скобки, пробелы, подчёркивания и обратные слэши
        formula = re.sub(r'[{}]', '', formula)
        formula = formula.replace(' ', '')
        formula = formula.replace('_', '')
        formula = formula.replace('\\', '')
        return formula
    clean_text = re.sub(r'\\\((.*?)\\\)', latex_to_text, clean_text)
    return clean_text

def format_output(parsed_data):
    formatted = ""
    for item in parsed_data:
        formatted += f"{item['question']}\n"
        
        # Выводим варианты ответов с буквами A) B) C) D)
        for idx, (label, answer) in enumerate(item['answers'], start=1):
            letter = chr(64 + idx)  # A=65, B=66, C=67, D=68
            formatted += f"{letter}) {answer}\n"
        
        # Собираем правильные ответы
        correct_answers = []
        for idx, (label, answer) in enumerate(item['answers'], start=1):
            if label == "+":
                letter = chr(64 + idx)
                correct_answers.append(letter)
        
        # Выводим правильные ответы
        if correct_answers:
            formatted += f"Ответы: {', '.join(correct_answers)}\n"
        
        formatted += "\n"
    return formatted

def save_to_txt(filename, content):
    with open(filename, "w", encoding="utf-8") as file:
        file.write(content)

def contains_images(parsed_data):
    for item in parsed_data:
        if extract_images_from_html(item['question']):
            return True
        for _, answer in item['answers']:
            if extract_images_from_html(answer):
                return True
    return False

def extract_image_filenames(text):
    """Извлекает имена файлов изображений из текста."""
    return re.findall(r'@(\S+)', text)

class ImgExtractor(HTMLParser):
    def __init__(self):
        super().__init__()
        self.imgs = []
    def handle_starttag(self, tag, attrs):
        if tag == 'img':
            attrs_dict = dict(attrs)
            src = attrs_dict.get('src', '')
            if src.startswith('@@PLUGINFILE@@/'):
                self.imgs.append(src.replace('@@PLUGINFILE@@/', ''))

def extract_images_from_html(html):
    parser = ImgExtractor()
    parser.feed(html)
    return parser.imgs

def save_images_from_xml(question_elem, temp_dir):
    """Сохраняет все изображения из <file> внутри вопроса, возвращает dict имя_файла: путь_к_файлу"""
    files = question_elem.findall('.//file')
    img_map = {}
    for file_elem in files:
        fname = file_elem.attrib.get('name')
        encoding = file_elem.attrib.get('encoding')
        data = file_elem.text
        if fname and encoding == 'base64' and data:
            img_path = os.path.join(temp_dir, fname)
            with open(img_path, 'wb') as f:
                f.write(base64.b64decode(data))
            img_map[fname] = img_path
    return img_map

def format_output_word(parsed_data, xml_filename):
    doc = Document()
    with tempfile.TemporaryDirectory() as temp_dir:
        for item in parsed_data:
            # Сохраняем изображения из <file> в temp_dir
            img_map = save_images_from_xml(item['question_elem'], temp_dir)
            # Парсим html вопроса
            html = item['question']
            parts = re.split(r'(<img[^>]+>)', html)
            for part in parts:
                if part.startswith('<img'):
                    imgs = extract_images_from_html(part)
                    for img in imgs:
                        # Декодируем имя файла из URL-формы
                        img_decoded = urllib.parse.unquote(img)
                        img_path = img_map.get(img_decoded)
                        if img_path and os.path.exists(img_path):
                            doc.add_picture(img_path, width=Inches(4))
                        else:
                            doc.add_paragraph(f'[Изображение {img} не найдено]')
                else:
                    text = clean_html(part).strip()
                    if text:
                        doc.add_paragraph(text)
            # Варианты ответов
            for idx, (label, answer) in enumerate(item['answers'], start=1):
                letter = chr(64 + idx)
                ans_html = answer
                ans_parts = re.split(r'(<img[^>]+>)', ans_html)
                ans_text = ''
                for ans_part in ans_parts:
                    if ans_part.startswith('<img'):
                        imgs = extract_images_from_html(ans_part)
                        for img in imgs:
                            img_decoded = urllib.parse.unquote(img)
                            img_path = img_map.get(img_decoded)
                            if img_path and os.path.exists(img_path):
                                doc.add_picture(img_path, width=Inches(4))
                            else:
                                doc.add_paragraph(f'[Изображение {img} не найдено]')
                    else:
                        ans_text += clean_html(ans_part).strip()
                doc.add_paragraph(f'{letter}) {ans_text}')
            # Правильные ответы
            correct_answers = []
            for idx, (label, answer) in enumerate(item['answers'], start=1):
                if label == '+':
                    letter = chr(64 + idx)
                    correct_answers.append(letter)
            if correct_answers:
                doc.add_paragraph(f'Ответы: {", ".join(correct_answers)}')
            doc.add_paragraph('')
    return doc

# Ищем XML файлы в текущей директории
xml_files = glob.glob("*.xml")

if not xml_files:
    print("XML файлы не найдены в текущей директории.")
    exit(1)

# Обрабатываем каждый найденный XML файл
for xml_file in xml_files:
    try:
        with open(xml_file, "r", encoding="utf-8") as file:
            xml_data = file.read()
        
        output_filename = os.path.splitext(xml_file)[0]
        parsed_data = extract_questions_and_answers(xml_data)
        
        if contains_images(parsed_data):
            doc = format_output_word(parsed_data, xml_file)
            doc.save(f"{output_filename}.docx")
            print(f"Данные из файла '{xml_file}' сохранены в файл '{output_filename}.docx'.")
        else:
            formatted_data = format_output(parsed_data)
            save_to_txt(f"{output_filename}.txt", formatted_data)
            print(f"Данные из файла '{xml_file}' сохранены в файл '{output_filename}.txt'.")
        
    except FileNotFoundError:
        print(f"Файл {xml_file} не найден.")
    except Exception as e:
        print(f"Ошибка при обработке файла {xml_file}: {e}")
